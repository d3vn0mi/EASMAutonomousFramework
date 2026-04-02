"""
Generate SoW and RoE documents from the existing DOCX templates.
Uses python-docx to find and replace placeholder markers.
"""
import os
from datetime import date
from io import BytesIO

from django.conf import settings
from docx import Document


PLACEHOLDERS = {
    "{{ENGAGEMENT_ID}}": lambda eng: eng.engagement_id,
    "{{CLIENT_NAME}}": lambda eng: eng.client.name,
    "{{ENGAGEMENT_NAME}}": lambda eng: eng.name,
    "{{TIER}}": lambda eng: eng.get_tier_display(),
    "{{START_DATE}}": lambda eng: eng.start_date.strftime("%d/%m/%Y") if eng.start_date else "TBD",
    "{{END_DATE}}": lambda eng: eng.end_date.strftime("%d/%m/%Y") if eng.end_date else "TBD",
    "{{DATE}}": lambda eng: date.today().strftime("%d/%m/%Y"),
    "{{PM_NAME}}": lambda eng: eng.project_manager.get_full_name() if eng.project_manager else "TBD",
}


def _template_path(template_name, language="en"):
    """Resolve the path to a DOCX template, accounting for language."""
    suffix = "_GR" if language == "el" else ""
    filename = f"{template_name}{suffix}.docx"
    return os.path.join(settings.RAVEN_TEMPLATES_DIR, filename)


def _replace_in_doc(doc, engagement):
    """Replace all known placeholders in paragraphs and tables."""
    for placeholder, value_fn in PLACEHOLDERS.items():
        value = value_fn(engagement)
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value))


def generate_sow(engagement):
    """Generate a Statement of Work document."""
    path = _template_path("RAVEN_SoW_Template", engagement.language)
    doc = Document(path)
    _replace_in_doc(doc, engagement)

    # Add scope items table
    scope_items = engagement.scope_items.filter(in_scope=True)
    if scope_items.exists():
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "{{SCOPE_TABLE}}" in cell.text:
                        cell.text = ""
                        for item in scope_items:
                            cell.add_paragraph(
                                f"{item.get_item_type_display()}: {item.value}"
                            )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_roe(engagement):
    """Generate a Rules of Engagement document."""
    path = _template_path("RAVEN_RoE_Template", engagement.language)
    doc = Document(path)
    _replace_in_doc(doc, engagement)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
