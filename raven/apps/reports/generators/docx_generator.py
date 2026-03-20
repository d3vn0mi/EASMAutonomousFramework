"""
Generate reports and documents from DOCX templates using python-docx.
Replaces placeholder markers in the templates with engagement data.
"""
import os
import logging
from datetime import datetime
from pathlib import Path

from django.conf import settings

logger = logging.getLogger(__name__)


def generate_document_from_template(template_path, engagement, doc_type):
    """
    Open a DOCX template, replace placeholders, and save to media.
    Returns the relative path to the generated file.
    """
    from docx import Document

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)

    # Build replacement map
    replacements = _build_replacements(engagement, doc_type)

    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)

    # Replace in headers/footers
    for section in doc.sections:
        for header in [section.header, section.footer]:
            if header:
                for paragraph in header.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)

    # Save
    output_dir = os.path.join(settings.MEDIA_ROOT, "documents", str(datetime.now().year))
    os.makedirs(output_dir, exist_ok=True)
    filename = f"{engagement.engagement_id}_{doc_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)

    # Return relative path for FileField
    return os.path.relpath(output_path, settings.MEDIA_ROOT)


def _build_replacements(engagement, doc_type):
    """Build a dictionary of placeholder -> value mappings."""
    client = engagement.client
    replacements = {
        "{{CLIENT_NAME}}": client.name,
        "{{ENGAGEMENT_ID}}": engagement.engagement_id,
        "{{ENGAGEMENT_NAME}}": engagement.name,
        "{{TIER}}": engagement.get_tier_display(),
        "{{START_DATE}}": str(engagement.start_date or "TBD"),
        "{{END_DATE}}": str(engagement.end_date or "TBD"),
        "{{DATE}}": datetime.now().strftime("%Y-%m-%d"),
        "{{PM_NAME}}": engagement.project_manager.get_full_name() if engagement.project_manager else "TBD",
        "[CLIENT NAME]": client.name,
        "[ENGAGEMENT ID]": engagement.engagement_id,
        "[DATE]": datetime.now().strftime("%Y-%m-%d"),
        "[DOMAIN]": "",
    }

    # Add scope domains
    domains = engagement.scope_items.filter(item_type="domain", in_scope=True)
    if domains.exists():
        replacements["[DOMAIN]"] = ", ".join(d.value for d in domains[:5])

    return replacements


def _replace_in_paragraph(paragraph, replacements):
    """Replace placeholders in a paragraph while preserving formatting."""
    full_text = paragraph.text
    for placeholder, value in replacements.items():
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, value)

    if full_text != paragraph.text:
        # Preserve formatting by updating runs
        if paragraph.runs:
            # Put all text in first run, clear others
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""


def generate_executive_report(engagement):
    """Generate executive report from tier-appropriate template."""
    lang_suffix = "_GR" if engagement.language == "el" else ""
    template_name = f"RAVEN_Report_{engagement.tier.title()}{lang_suffix}.docx"
    template_path = os.path.join(settings.RAVEN_TEMPLATES_DIR, template_name)
    return generate_document_from_template(template_path, engagement, "executive")


def generate_technical_report(engagement):
    """Generate technical report from tier-appropriate template."""
    lang_suffix = "_GR" if engagement.language == "el" else ""
    template_name = f"RAVEN_Report_{engagement.tier.title()}{lang_suffix}.docx"
    template_path = os.path.join(settings.RAVEN_TEMPLATES_DIR, template_name)
    return generate_document_from_template(template_path, engagement, "technical")
